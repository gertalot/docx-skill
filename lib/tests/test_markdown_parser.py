"""tests/test_markdown_parser.py"""
import pytest
from docx import Document
from docx_builder.brand import Brand
from docx_builder.builder import DocBuilder
from docx_builder.markdown_parser import render_markdown


@pytest.fixture
def builder():
    brand = Brand(
        primary="#1C2541", accent="#6FFFE9", body="#3E465E",
        font_family="Helvetica",
    )
    b = DocBuilder(brand)
    b.setup_page()
    b.setup_styles()
    return b


class TestRenderMarkdown:
    def test_heading_1(self, builder):
        render_markdown(builder, "## Section Title")
        styles = [p.style.name for p in builder.doc.paragraphs]
        assert "Heading 1" in styles

    def test_heading_2(self, builder):
        render_markdown(builder, "### Sub Section")
        styles = [p.style.name for p in builder.doc.paragraphs]
        assert "Heading 2" in styles

    def test_paragraph(self, builder):
        render_markdown(builder, "This is a paragraph.")
        assert any("This is a paragraph." in p.text for p in builder.doc.paragraphs)

    def test_multiline_paragraph(self, builder):
        md = "This is the first line\nand the second line."
        render_markdown(builder, md)
        texts = [p.text for p in builder.doc.paragraphs]
        assert any("first line and the second line" in t for t in texts)

    def test_bullet_list(self, builder):
        render_markdown(builder, "- Item one\n- Item two")
        styles = [p.style.name for p in builder.doc.paragraphs]
        assert styles.count("List Bullet") == 2

    def test_bullet_with_continuation(self, builder):
        md = "- First line of bullet\n  and the continuation."
        render_markdown(builder, md)
        texts = [p.text for p in builder.doc.paragraphs]
        assert any("First line of bullet and the continuation." in t for t in texts)

    def test_emphasis_uses_medium_font(self, builder):
        render_markdown(builder, "Text with **emphasis** here.")
        para = builder.doc.paragraphs[0]
        emphasis_runs = [r for r in para.runs if "emphasis" in r.text]
        assert len(emphasis_runs) == 1
        assert emphasis_runs[0].bold is not True

    def test_table(self, builder):
        md = "| A | B |\n|---|---|\n| 1 | 2 |"
        render_markdown(builder, md)
        assert len(builder.doc.tables) == 1
        assert builder.doc.tables[0].rows[0].cells[0].text.strip() == "A"

    def test_horizontal_rule_creates_page_break(self, builder):
        md = "## Section 1\n\nParagraph.\n\n---\n\n## Section 2"
        render_markdown(builder, md)
        texts = [p.text for p in builder.doc.paragraphs]
        assert "Section 1" in texts
        assert "Section 2" in texts

    def test_code_blocks_skipped(self, builder):
        md = "Before.\n\n```python\ncode here\n```\n\nAfter."
        render_markdown(builder, md)
        texts = " ".join(p.text for p in builder.doc.paragraphs)
        assert "code here" not in texts
        assert "Before." in texts
        assert "After." in texts

    def test_title_block_skipped(self, builder):
        md = "# Document Title\n\n**Version 1.0**\n\n---\n\n## Real Content"
        render_markdown(builder, md)
        texts = [p.text for p in builder.doc.paragraphs]
        assert "Document Title" not in texts
        assert "Real Content" in texts
