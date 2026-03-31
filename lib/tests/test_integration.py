"""tests/test_integration.py — full pipeline test."""
from pathlib import Path

import pytest
from docx import Document

from docx_builder import Brand, DocBuilder
from docx_builder.markdown_parser import render_markdown


QRTR_WHITEPAPER = Path.home() / "Code/Qrtr.ai/qrtr-platform/docs/security-whitepaper.md"
QRTR_LOGO = Path.home() / "Code/Qrtr.ai/qrtr-platform/Qrtr.Frontend/src/assets/logo-dark.svg"


@pytest.mark.skipif(
    not QRTR_WHITEPAPER.exists(), reason="Qrtr whitepaper not found"
)
class TestQrtrWhitepaper:
    def test_full_pipeline(self, tmp_path):
        brand = Brand(
            primary="#1C2541",
            accent="#6FFFE9",
            body="#3E465E",
            font_family="Avenir Next",
            emphasis_weight=500,
            heading_weight=700,
            logo_path=QRTR_LOGO if QRTR_LOGO.exists() else None,
        )

        builder = DocBuilder(brand)
        builder.setup_page(
            size="A4",
            margins={"left": 3.0, "right": 3.0, "top": 3.5, "bottom": 2.54},
        )
        builder.setup_styles(language="en-GB", justify=True)
        builder.add_cover(
            title="Security Whitepaper",
            subtitle="Version 1.0 \u2014 March 2026",
            confidential=True,
        )
        builder.add_toc()
        builder.setup_header(logo=True, title="Qrtr Security Whitepaper")
        builder.setup_footer(center="page_number", right="Confidential")
        render_markdown(builder, QRTR_WHITEPAPER)

        output = tmp_path / "whitepaper.docx"
        builder.save(output)

        # Validate output
        doc = Document(str(output))
        assert len(doc.paragraphs) > 50
        assert len(doc.tables) >= 2

        # Check headings exist
        headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
        assert "1. Executive Summary" in headings
        assert "8. AI & LLM Security" in headings
