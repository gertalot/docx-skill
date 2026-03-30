"""tests/test_builder.py"""
from pathlib import Path

import pytest
from docx import Document

from docx_builder.brand import Brand
from docx_builder.builder import DocBuilder


@pytest.fixture
def brand():
    return Brand(
        primary="#1C2541",
        accent="#6FFFE9",
        body="#3E465E",
        font_family="Helvetica",
    )


class TestDocBuilder:
    def test_creates_document(self, brand):
        builder = DocBuilder(brand)
        assert builder.doc is not None

    def test_setup_page_a4(self, brand):
        builder = DocBuilder(brand)
        builder.setup_page(size="A4")
        section = builder.doc.sections[0]
        # Cm(21.0) = 7560000 EMU, but python-docx round-trips through twips
        assert section.page_width == pytest.approx(7560000, abs=1000)

    def test_setup_page_letter(self, brand):
        builder = DocBuilder(brand)
        builder.setup_page(size="Letter")
        section = builder.doc.sections[0]
        assert section.page_width == 7772400  # 8.5 inches in EMU

    def test_setup_page_custom_margins(self, brand):
        builder = DocBuilder(brand)
        builder.setup_page(margins={"left": 2.0, "right": 2.0, "top": 3.0, "bottom": 2.0})
        section = builder.doc.sections[0]
        # Cm(2.0) = 720000 EMU, but python-docx round-trips through twips
        assert section.left_margin == pytest.approx(720000, abs=1000)

    def test_setup_styles_sets_font(self, brand):
        builder = DocBuilder(brand)
        builder.setup_page()
        builder.setup_styles()
        normal = builder.doc.styles["Normal"]
        assert normal.font.name is not None

    def test_add_cover(self, brand):
        builder = DocBuilder(brand)
        builder.setup_page()
        builder.setup_styles()
        builder.add_cover(title="Test Document", subtitle="v1.0")
        assert len(builder.doc.paragraphs) > 0

    def test_setup_header_with_title(self, brand):
        builder = DocBuilder(brand)
        builder.setup_page()
        builder.setup_styles()
        builder.setup_header(title="My Document")
        section = builder.doc.sections[0]
        assert section.different_first_page_header_footer is True

    def test_setup_footer_page_number(self, brand):
        builder = DocBuilder(brand)
        builder.setup_page()
        builder.setup_styles()
        builder.setup_footer(center="page_number", right="Confidential")
        section = builder.doc.sections[0]
        assert section.footer is not None

    def test_save(self, brand, tmp_path):
        builder = DocBuilder(brand)
        builder.setup_page()
        builder.setup_styles()
        builder.add_cover(title="Test")
        path = tmp_path / "test.docx"
        builder.save(path)
        assert path.exists()
        doc = Document(str(path))
        assert len(doc.paragraphs) > 0

    def test_add_toc(self, brand):
        builder = DocBuilder(brand)
        builder.setup_page()
        builder.setup_styles()
        builder.add_toc()
        assert len(builder.doc.paragraphs) >= 2
