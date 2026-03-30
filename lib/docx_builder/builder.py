"""DocBuilder — branded Word document builder wrapping python-docx.

Extracts and generalises the proven approach from the Qrtr whitepaper
prototype, replacing hardcoded colours/fonts with Brand properties.
"""
from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Pt

if TYPE_CHECKING:
    from docx_builder.brand import Brand

# Page size dimensions in EMU (English Metric Units)
_PAGE_SIZES = {
    "A4": (Cm(21.0), Cm(29.7)),
    "Letter": (Inches(8.5), Inches(11)),
}

_DEFAULT_MARGINS = {"left": 3.0, "right": 3.0, "top": 3.5, "bottom": 2.54}


# ── Low-level helpers ────────────────────────────────────────────


def _accent_rule(para, accent_hex: str) -> None:
    """Add accent-coloured bottom-border to a paragraph."""
    para._p.get_or_add_pPr().append(
        parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="6" w:space="2" '
            f'w:color="{accent_hex}"/>'
            f"</w:pBdr>"
        )
    )


def _bottom_border(para, color_hex: str, sz: int = 4, space: int = 4) -> None:
    """Add a bottom border to a paragraph."""
    para._p.get_or_add_pPr().append(
        parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="{sz}" w:space="{space}" '
            f'w:color="{color_hex}"/>'
            f"</w:pBdr>"
        )
    )


def _top_border(para, color_hex: str, sz: int = 4, space: int = 0) -> None:
    """Add a top border to a paragraph."""
    para._p.get_or_add_pPr().append(
        parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="{sz}" w:space="{space}" '
            f'w:color="{color_hex}"/>'
            f"</w:pBdr>"
        )
    )


def _remove_table_borders(tbl) -> None:
    """Remove all borders from a table."""
    tbl._tbl.tblPr.append(
        parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '<w:top w:val="none" w:sz="0" w:space="0"/>'
            '<w:left w:val="none" w:sz="0" w:space="0"/>'
            '<w:bottom w:val="none" w:sz="0" w:space="0"/>'
            '<w:right w:val="none" w:sz="0" w:space="0"/>'
            '<w:insideH w:val="none" w:sz="0" w:space="0"/>'
            '<w:insideV w:val="none" w:sz="0" w:space="0"/>'
            "</w:tblBorders>"
        )
    )


def _add_field(para, field_code: str) -> None:
    """Insert a Word field code (e.g. PAGE, NUMPAGES) into a paragraph."""
    run = para.add_run()
    el = OxmlElement("w:fldChar")
    el.set(qn("w:fldCharType"), "begin")
    run._r.append(el)

    run = para.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    run._r.append(instr)

    run = para.add_run()
    el = OxmlElement("w:fldChar")
    el.set(qn("w:fldCharType"), "separate")
    run._r.append(el)


def _end_field(para) -> None:
    """Close a Word field code in a paragraph."""
    run = para.add_run()
    el = OxmlElement("w:fldChar")
    el.set(qn("w:fldCharType"), "end")
    run._r.append(el)


def _page_number_field(para, brand: Brand) -> None:
    """Insert a PAGE field code into a paragraph with brand styling."""
    _add_field(para, "PAGE")

    # Placeholder text (replaced when Word updates the field)
    run = para.add_run("1")
    run.font.name = brand.font_regular
    run.font.size = Pt(9)
    run.font.color.rgb = brand.body_rgb.to_docx()

    _end_field(para)


def _page_x_of_y_field(para, brand: Brand) -> None:
    """Insert 'Page X of Y' field codes with brand styling."""
    # "Page " prefix
    run = para.add_run("Page ")
    run.font.name = brand.font_regular
    run.font.size = Pt(9)
    run.font.color.rgb = brand.body_rgb.to_docx()

    # PAGE field
    _add_field(para, "PAGE")
    run = para.add_run("1")
    run.font.name = brand.font_regular
    run.font.size = Pt(9)
    run.font.color.rgb = brand.body_rgb.to_docx()
    _end_field(para)

    # " of " separator
    run = para.add_run(" of ")
    run.font.name = brand.font_regular
    run.font.size = Pt(9)
    run.font.color.rgb = brand.body_rgb.to_docx()

    # NUMPAGES field
    _add_field(para, "NUMPAGES")
    run = para.add_run("1")
    run.font.name = brand.font_regular
    run.font.size = Pt(9)
    run.font.color.rgb = brand.body_rgb.to_docx()
    _end_field(para)


def _add_tab_stops(para_format, text_width_emu: int) -> None:
    """Add centre and right-aligned tab stops to a paragraph format."""
    tab_stops = para_format.tab_stops
    tab_stops.add_tab_stop(text_width_emu // 2, WD_TAB_ALIGNMENT.CENTER)
    tab_stops.add_tab_stop(text_width_emu, WD_TAB_ALIGNMENT.RIGHT)


# ── DocBuilder ───────────────────────────────────────────────────


class DocBuilder:
    """Branded Word document builder.

    Wraps python-docx's Document with methods for cover pages, headers,
    footers, table of contents, and styled content. Uses Brand for all
    colours, fonts, and logo.
    """

    def __init__(self, brand: Brand) -> None:
        self.brand = brand
        self.doc = Document()

    # ── Page setup ───────────────────────────────────────────────

    def setup_page(
        self,
        size: str = "A4",
        margins: dict[str, float] | None = None,
    ) -> None:
        """Configure page size and margins.

        Args:
            size: "A4" or "Letter".
            margins: dict with left/right/top/bottom in cm.
                     Defaults to 3cm left/right, 3.5cm top, 2.54cm bottom.
        """
        if size not in _PAGE_SIZES:
            raise ValueError(
                f"Unknown page size {size!r}. Supported: {', '.join(_PAGE_SIZES)}"
            )

        section = self.doc.sections[0]
        width, height = _PAGE_SIZES[size]
        section.page_width = width
        section.page_height = height

        m = {**_DEFAULT_MARGINS, **(margins or {})}
        section.left_margin = Cm(m["left"])
        section.right_margin = Cm(m["right"])
        section.top_margin = Cm(m["top"])
        section.bottom_margin = Cm(m["bottom"])

    # ── Styles ───────────────────────────────────────────────────

    def setup_styles(
        self,
        language: str = "en-GB",
        justify: bool = True,
    ) -> None:
        """Configure document styles using brand fonts and colours.

        Args:
            language: Language code for spell checking (e.g. "en-GB").
            justify: Whether to justify normal text.
        """
        brand = self.brand

        # Normal style
        normal = self.doc.styles["Normal"]
        normal.font.name = brand.font_regular
        normal.font.size = Pt(11)
        normal.font.color.rgb = brand.body_rgb.to_docx()
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal.paragraph_format.line_spacing = 1.3
        if justify:
            normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Language for spell-check
        lang = OxmlElement("w:lang")
        lang.set(qn("w:val"), language)
        normal.element.get_or_add_rPr().append(lang)

        # Heading styles
        for name, size, before, after in [
            ("Heading 1", 20, 0, 12),
            ("Heading 2", 16, 18, 8),
        ]:
            h = self.doc.styles[name]
            h.font.name = brand.font_heading
            h.font.size = Pt(size)
            h.font.color.rgb = brand.primary_rgb.to_docx()
            h.paragraph_format.space_before = Pt(before)
            h.paragraph_format.space_after = Pt(after)
            h.paragraph_format.keep_with_next = True

        # List Bullet styles
        for name in ("List Bullet", "List Bullet 2"):
            try:
                lb = self.doc.styles[name]
                lb.font.name = brand.font_regular
                lb.font.size = Pt(11)
                lb.font.color.rgb = brand.body_rgb.to_docx()
                lb.paragraph_format.space_after = Pt(4)
                lb.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                lb.paragraph_format.line_spacing = 1.3
            except KeyError:
                pass

    # ── Cover page ───────────────────────────────────────────────

    def add_cover(
        self,
        title: str,
        subtitle: str | None = None,
        confidential: bool = False,
    ) -> None:
        """Add a cover page with title, optional subtitle, and optional logo.

        Args:
            title: Main document title.
            subtitle: Version/date line below the title.
            confidential: Whether to show "Confidential" at the bottom.
        """
        brand = self.brand

        # Logo (if available)
        logo = brand.logo_png(width_px=300)
        if logo is not None:
            p = self.doc.add_paragraph()
            p.add_run().add_picture(logo, width=Inches(1.8))
            p.paragraph_format.space_after = Pt(150)
        else:
            # Spacer when no logo
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(200)

        # Title
        p = self.doc.add_paragraph()
        r = p.add_run(title)
        r.font.name = brand.font_heading
        r.font.size = Pt(32)
        r.font.color.rgb = brand.primary_rgb.to_docx()
        p.paragraph_format.space_after = Pt(4)

        # Accent rule under title
        p = self.doc.add_paragraph()
        _accent_rule(p, brand.accent_hex)
        p.paragraph_format.space_after = Pt(8)

        # Subtitle
        if subtitle is not None:
            p = self.doc.add_paragraph()
            r = p.add_run(subtitle)
            r.font.name = brand.font_regular
            r.font.size = Pt(14)
            r.font.color.rgb = brand.body_rgb.to_docx()

        # Confidential marker
        if confidential:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(280)
            r = p.add_run("Confidential")
            r.font.name = brand.font_regular
            r.font.size = Pt(10)
            r.font.color.rgb = brand.body_rgb.to_docx()
            r.italic = True

        self.doc.add_page_break()

    # ── Table of Contents ────────────────────────────────────────

    def add_toc(
        self,
        title: str = "Table of Contents",
        levels: int = 2,
    ) -> None:
        """Insert a Table of Contents field.

        Args:
            title: Title shown above the TOC.
            levels: Number of heading levels to include (1-9).
        """
        brand = self.brand

        # TOC title
        p = self.doc.add_paragraph()
        r = p.add_run(title)
        r.font.name = brand.font_heading
        r.font.size = Pt(20)
        r.font.color.rgb = brand.primary_rgb.to_docx()
        p.paragraph_format.space_after = Pt(24)

        # TOC field code
        p = self.doc.add_paragraph()

        _add_field(p, f'TOC \\o "1-{levels}" \\h \\z \\u')

        # Placeholder text
        run = p.add_run(
            "Open in Word and update this field to populate (Ctrl+A, F9)"
        )
        run.font.name = brand.font_regular
        run.font.size = Pt(11)
        run.font.color.rgb = brand.body_rgb.to_docx()
        run.italic = True

        _end_field(p)

        self.doc.add_page_break()

    # ── Header ───────────────────────────────────────────────────

    def setup_header(
        self,
        logo: bool = True,
        title: str | None = None,
        separator: bool = True,
    ) -> None:
        """Configure the document header.

        When logo=True and brand has a logo, uses a borderless table to
        place the logo left and title right (tab stops cannot vertically
        align inline images).

        When logo=False or no logo is available, uses tab stops for a
        text-only layout.

        Args:
            logo: Whether to include the brand logo in the header.
            title: Text to show in the header (right-aligned).
            separator: Whether to add a border line below the header.
        """
        brand = self.brand
        section = self.doc.sections[0]
        section.different_first_page_header_footer = True
        text_width = section.page_width - section.left_margin - section.right_margin

        header = section.header
        header.is_linked_to_previous = False

        # Remove default paragraph(s)
        for p in header.paragraphs:
            p._p.getparent().remove(p._p)

        logo_png = brand.logo_png(width_px=120) if logo else None

        if logo_png is not None:
            # Table layout for logo + title
            tbl = header.add_table(1, 2, text_width)
            tbl.autofit = False
            _remove_table_borders(tbl)

            # Left cell: logo
            lp = tbl.cell(0, 0).paragraphs[0]
            lp.add_run().add_picture(logo_png, width=Inches(0.7))

            # Right cell: title
            if title is not None:
                rp = tbl.cell(0, 1).paragraphs[0]
                rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r = rp.add_run(title)
                r.font.name = brand.font_regular
                r.font.size = Pt(9)
                r.font.color.rgb = brand.body_rgb.to_docx()
        elif title is not None:
            # Text-only layout with right tab stop
            p = header.add_paragraph()
            pf = p.paragraph_format
            pf.tab_stops.add_tab_stop(text_width, WD_TAB_ALIGNMENT.RIGHT)

            run = p.add_run()
            run.add_tab()
            run = p.add_run(title)
            run.font.name = brand.font_regular
            run.font.size = Pt(9)
            run.font.color.rgb = brand.body_rgb.to_docx()

        # Separator line
        if separator:
            sep = header.add_paragraph()
            sep.paragraph_format.space_before = Pt(4)
            sep.paragraph_format.space_after = Pt(0)
            _top_border(sep, brand.border_hex)

    # ── Footer ───────────────────────────────────────────────────

    def setup_footer(
        self,
        left: str | None = None,
        center: str | None = None,
        right: str | None = None,
    ) -> None:
        """Configure the document footer using tab stops.

        Each zone (left, center, right) can be:
        - A string literal to display as text
        - "page_number" for a PAGE field
        - "page_x_of_y" for a "Page X of Y" field
        - None to leave that zone empty

        Args:
            left: Content for the left zone.
            center: Content for the centre zone.
            right: Content for the right zone.
        """
        brand = self.brand
        section = self.doc.sections[0]
        text_width = section.page_width - section.left_margin - section.right_margin

        footer = section.footer
        footer.is_linked_to_previous = False

        # Remove default paragraph(s)
        for p in footer.paragraphs:
            p._p.getparent().remove(p._p)

        p = footer.add_paragraph()
        _add_tab_stops(p.paragraph_format, text_width)

        zones = [left, center, right]
        for i, zone in enumerate(zones):
            if i > 0:
                run = p.add_run()
                run.add_tab()

            if zone is None:
                continue
            elif zone == "page_number":
                _page_number_field(p, brand)
            elif zone == "page_x_of_y":
                _page_x_of_y_field(p, brand)
            else:
                run = p.add_run(zone)
                run.font.name = brand.font_regular
                run.font.size = Pt(9)
                run.font.color.rgb = brand.body_rgb.to_docx()
                if zone.lower() == "confidential":
                    run.italic = True

    # ── Save ─────────────────────────────────────────────────────

    def save(self, path: str | Path) -> None:
        """Save the document to the specified path.

        Args:
            path: File path for the output .docx file.
        """
        self.doc.save(str(path))
