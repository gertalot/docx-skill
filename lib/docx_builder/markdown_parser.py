"""Markdown-to-docx renderer using DocBuilder and Brand.

Extracts and generalises the render_body / render_table / add_runs logic
from the Qrtr whitepaper prototype, replacing hardcoded colours and fonts
with Brand properties.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import TYPE_CHECKING

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Pt, RGBColor as DocxRGBColor

from docx_builder.builder import _accent_rule

if TYPE_CHECKING:
    from docx.document import Document
    from docx.table import _Cell
    from docx.text.paragraph import Paragraph

    from docx_builder.brand import Brand
    from docx_builder.builder import DocBuilder

WHITE = DocxRGBColor(0xFF, 0xFF, 0xFF)


# ── Low-level cell helpers ────────────────────────────────────────


def _cell_shading(cell: _Cell, hex_color: str) -> None:
    """Apply background shading to a table cell."""
    cell._tc.get_or_add_tcPr().append(
        parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
        )
    )


def _cell_borders(cell: _Cell, color: str, sz: int = 4) -> None:
    """Apply uniform borders to a table cell."""
    cell._tc.get_or_add_tcPr().append(
        parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            + "".join(
                f'<w:{side} w:val="single" w:sz="{sz}" w:space="0" '
                f'w:color="{color}"/>'
                for side in ("top", "left", "bottom", "right")
            )
            + "</w:tcBorders>"
        )
    )


# ── Inline text parsing ──────────────────────────────────────────


def add_runs(para: Paragraph, text: str, brand: Brand) -> None:
    """Parse ``**emphasis**`` markers and add formatted runs.

    Emphasis text uses ``brand.font_emphasis`` (a heavier font weight)
    rather than ``run.bold = True``, giving proper typographic emphasis
    that respects the brand's font family.
    """
    for part in re.split(r"(\*\*.*?\*\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = para.add_run(part[2:-2])
            r.font.name = brand.font_emphasis
        else:
            para.add_run(part)


# ── Table rendering ──────────────────────────────────────────────


def render_table(builder: DocBuilder, raw_rows: list[str]) -> None:
    """Render a markdown table with branded styling.

    Header row gets the primary colour background with heading font.
    Data rows alternate with a light background. All cells get brand
    border colour.
    """
    brand = builder.brand
    doc = builder.doc

    def parse_row(line: str) -> list[str]:
        return [c.strip() for c in line.strip().strip("|").split("|")]

    header = parse_row(raw_rows[0])
    data = [parse_row(r) for r in raw_rows[2:]]  # skip separator row
    cols = len(header)

    tbl = doc.add_table(rows=1 + len(data), cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True

    # Header row
    for ci, text in enumerate(header):
        cell = tbl.rows[0].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.name = brand.font_heading
        r.font.size = Pt(10)
        r.font.color.rgb = WHITE
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        _cell_shading(cell, brand.primary_hex)
        _cell_borders(cell, brand.border_hex)

    # Data rows
    for ri, row in enumerate(data):
        for ci in range(min(len(row), cols)):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, row[ci], brand)
            for run in p.runs:
                run.font.name = brand.font_regular
                run.font.size = Pt(10)
                run.font.color.rgb = brand.body_rgb.to_docx()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            _cell_borders(cell, brand.border_hex)
            if ri % 2 == 1:
                _cell_shading(cell, brand.light_bg_hex)

    # Space after table
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ── Body rendering ───────────────────────────────────────────────


def _has_title_block(lines: list[str]) -> bool:
    """Detect whether the markdown has a title block to skip.

    A title block is present when the first non-blank line is a level-1
    heading (``# Title``) and a ``---`` separator appears before any
    body content (``##`` heading).
    """
    first_content = None
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if first_content is None:
            first_content = stripped
            if not stripped.startswith("# "):
                return False
            continue
        if stripped == "---":
            return True
        if stripped.startswith("## "):
            return False
    return False


def render_markdown(builder: DocBuilder, md: str | Path) -> None:
    """Render markdown content into a DocBuilder document.

    Handles headings, paragraphs, bullets, tables, code blocks,
    horizontal rules, and title blocks.

    Args:
        builder: DocBuilder instance with page and styles already set up.
        md: Markdown text as a string, or a Path to a markdown file.
    """
    if isinstance(md, Path):
        md = md.read_text(encoding="utf-8")

    brand = builder.brand
    doc = builder.doc
    lines = md.split("\n")
    n = len(lines)
    i = 0
    in_code = False

    # Detect whether a title block exists: content before the first ---
    # that starts with a level-1 heading (# Title).
    past_title = not _has_title_block(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Code blocks: toggle and skip
        if stripped.startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            i += 1
            continue

        # Skip title block (everything before the first ---)
        if not past_title:
            if stripped == "---":
                past_title = True
            i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Horizontal rule -> page break
        if stripped == "---":
            doc.add_page_break()
            i += 1
            continue

        # ## -> Heading 1 with accent underline
        if line.startswith("## "):
            p = doc.add_paragraph(line[3:].strip(), style="Heading 1")
            _accent_rule(p, brand.accent_hex)
            i += 1
            continue

        # ### -> Heading 2
        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            i += 1
            continue

        # Table
        if stripped.startswith("|") and "|" in stripped[1:]:
            tbl_lines: list[str] = []
            while i < n and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            if len(tbl_lines) >= 2:
                render_table(builder, tbl_lines)
            continue

        # Bullet list item
        if stripped.startswith("- "):
            indent = len(line) - len(line.lstrip())
            text = stripped[2:]
            # Gather indented continuation lines
            while i + 1 < n:
                nxt = lines[i + 1]
                ns = nxt.strip()
                if not ns or ns.startswith("- "):
                    break
                if not (nxt.startswith("  ") or nxt.startswith("\t")):
                    break
                text += " " + ns
                i += 1
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Cm(1.27)
            add_runs(p, text, brand)
            i += 1
            continue

        # Regular paragraph: gather soft-wrapped continuation lines
        text = stripped
        while i + 1 < n:
            nxt = lines[i + 1]
            ns = nxt.strip()
            if (
                not ns
                or nxt.startswith("  ")
                or nxt.startswith("\t")
                or ns.startswith("#")
                or ns.startswith("- ")
                or ns == "---"
                or ns.startswith("```")
                or (ns.startswith("|") and "|" in ns[1:])
            ):
                break
            text += " " + ns
            i += 1
        p = doc.add_paragraph(style="Normal")
        add_runs(p, text, brand)
        i += 1
